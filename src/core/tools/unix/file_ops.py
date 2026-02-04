"""
File Operations Tools
Essential file system operations with workspace isolation
"""
import logging
import subprocess
from typing import Dict, List, Any
from pathlib import Path
import os
import time

from ..base import resolve_path, execute_command, DEFAULT_WORKSPACE

logger = logging.getLogger(__name__)


def read_file(file_path: str, workspace: Path = DEFAULT_WORKSPACE,
              start_line: int = None, limit_lines: int = None, 
              auto_truncate: bool = True, _context: dict = None) -> Dict[str, Any]:
    """Read file content with smart pagination for any file type"""
    try:
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ File not found: '{file_path}' - check if the file exists or path is correct"}
        
        # Check file size (limit to 10MB)
        file_size = resolved_path.stat().st_size
        if file_size > 10 * 1024 * 1024:
            return {'success': False, 'output': f"❌ File too large: '{file_path}' ({file_size // (1024*1024)}MB) - files over 10MB not supported"}
        
        # Detect file type and extract content
        file_type = _detect_file_type(resolved_path)
        content = _extract_content_to_text(resolved_path, file_type, workspace)
        
        if not content.strip():
            return {'success': False, 'output': f"❌ No content extracted from '{file_path}' - file may be empty or unsupported format"}
        
        # Apply smart pagination
        lines = content.split('\n')
        total_lines = len(lines)
        
        # Determine start and end positions
        if start_line is None:
            start_pos = 0
        else:
            start_pos = max(0, start_line - 1)  # Convert to 0-indexed
        
        # Auto-truncation based on actual token count with line-based cuts
        logger.debug(f"read_file - auto_truncate={auto_truncate}, limit_lines={limit_lines}, has_context={_context is not None}")
        if auto_truncate and limit_lines is None:
            try:
                # Try to use tiktoken for accurate token counting if available
                try:
                    import tiktoken
                    enc = tiktoken.get_encoding("cl100k_base")
                    actual_tokens = len(enc.encode(content))
                    token_manager = True  # Flag that we have accurate counting
                except ImportError:
                    # Fallback to character-based estimation
                    actual_tokens = len(content) // 4
                    token_manager = None
                
                # Use context-aware thresholds if available
                if _context:
                    available_tokens = _context['available_tokens']
                    threshold = available_tokens
                    # Use only 60% of available space to account for conversation growth
                    target_tokens = min(int(available_tokens * 0.6), actual_tokens)
                    logger.debug(f" Context available - file: {actual_tokens:,}, threshold: {threshold:,}, target: {target_tokens:,}")
                else:
                    # Fallback to conservative static limits
                    threshold = 15000
                    target_tokens = 10000
                    logger.debug(f" No context - file: {actual_tokens:,}, threshold: {threshold:,}, target: {target_tokens:,}")
                
                # Truncate if content exceeds available space
                if actual_tokens > threshold:
                    logger.debug(f" TRUNCATION TRIGGERED! {actual_tokens:,} > {threshold:,}")
                    # Calculate average tokens per line for intelligent truncation
                    tokens_per_line = actual_tokens / total_lines if total_lines > 0 else 1
                    # Target available space or reasonable portion
                    safe_lines = max(100, int(target_tokens / tokens_per_line))
                    
                    end_pos = min(start_pos + safe_lines, total_lines)
                    truncated = True
                    
                    continuation_hint = f"\n\n📄 Document truncated - showing lines {start_pos + 1}-{end_pos} of {total_lines}."
                    continuation_hint += f"\nActual tokens: {actual_tokens:,} → truncated to ~{int(safe_lines * tokens_per_line):,} tokens."
                    continuation_hint += f"\nTo read more: read_file('{file_path}', start_line={end_pos + 1}, limit_lines={safe_lines})"
                else:
                    end_pos = total_lines
                    truncated = False
                    continuation_hint = ""
                    
            except Exception as e:
                # Fallback to character estimation if TokenManager fails
                logger.warning(f"Token counting failed, using estimation: {e}")
                estimated_tokens = len(content) // 4
                if estimated_tokens > 15000:
                    # Conservative fallback - limit to 500 lines
                    max_lines = 500
                    end_pos = min(start_pos + max_lines, total_lines)
                    truncated = True
                    continuation_hint = f"\n\n📄 Document truncated - showing lines {start_pos + 1}-{end_pos} of {total_lines}."
                    continuation_hint += f"\nEstimated {estimated_tokens:,} tokens (tokenizer unavailable). To read more: read_file('{file_path}', start_line={end_pos + 1}, limit_lines=500)"
                else:
                    end_pos = total_lines
                    truncated = False
                    continuation_hint = ""
        else:
            if limit_lines is None:
                end_pos = total_lines
            else:
                end_pos = min(start_pos + limit_lines, total_lines)
            truncated = False
            continuation_hint = ""
        
        # Extract the specified lines
        selected_lines = lines[start_pos:end_pos]
        result_content = '\n'.join(selected_lines)
        
        # Build output
        # Always show range info to make truncation clear
        range_info = f" (lines {start_pos + 1}-{end_pos} of {total_lines})"
        output = f"📄 {file_path}{range_info}:\n{result_content}"
        if continuation_hint:
            output += continuation_hint
        
        # Calculate final token info for transparency
        try:
            import tiktoken
            enc = tiktoken.get_encoding("cl100k_base")
            final_tokens = len(enc.encode(result_content))
        except ImportError:
            final_tokens = len(result_content) // 4
        
        # Add token transparency if context is available
        if _context:
            token_usage_info = f"\n\n🧠 Token Usage:\n"
            token_usage_info += f"• File content: {final_tokens:,} tokens\n"
            token_usage_info += f"• Context used: {_context['current_tokens']:,} / {_context['max_tokens']:,} tokens\n"
            token_usage_info += f"• Available space: {_context['available_tokens']:,} tokens remaining\n"
            if truncated:
                token_usage_info += f"• ⚠️  Content truncated to fit available space"
            
            output += token_usage_info
        
        result = {
            'success': True,
            'output': output,
            'content': result_content,
            'file_type': file_type,
            'lines_shown': len(selected_lines),
            'total_lines': total_lines,
            'line_range': f"{start_pos + 1}-{end_pos}",
            'truncated': truncated,
            'file_size': file_size
        }
        
        # Add structured token info for programmatic access
        if _context:
            result['token_info'] = {
                'file_tokens': final_tokens,
                'context_used': _context['current_tokens'],
                'available_tokens': _context['available_tokens'],
                'max_tokens': _context['max_tokens'],
                'truncated': truncated
            }
        
        return result
        
    except UnicodeDecodeError:
        return {'success': False, 'output': f"❌ Cannot read '{file_path}' - file appears to be binary, not text"}
    except PermissionError:
        return {'success': False, 'output': f"❌ Permission denied reading '{file_path}' - insufficient privileges to access this file"}
    except IsADirectoryError:
        return {'success': False, 'output': f"❌ Cannot read '{file_path}' - this is a directory, not a file"}
    except OSError as e:
        if "No such file or directory" in str(e):
            return {'success': False, 'output': f"❌ File not found: '{file_path}' - path does not exist"}
        else:
            return {'success': False, 'output': f"❌ System error reading '{file_path}': {str(e)} - file system issue"}
    except Exception as e:
        return {'success': False, 'output': f"❌ Error reading '{file_path}': {str(e)} - unable to read file content"}


def _detect_file_type(file_path: Path) -> str:
    """Detect file type based on extension"""
    extension = file_path.suffix.lower()
    if extension == '.pdf':
        return 'pdf'
    elif extension == '.docx':
        return 'docx'
    elif extension in ['.xlsx', '.xls']:
        return 'excel'
    elif extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.log', 
                       '.csv', '.c', '.cpp', '.h', '.hpp', '.php', '.rb', '.go', '.rs', '.sh', '.bat', '.ps1',
                       '.sql', '.r', '.m', '.swift', '.kt', '.scala', '.pl', '.lua', '.vim', '.ini', '.conf']:
        return 'text'
    else:
        return 'text'  # Default to text for unknown extensions


def _extract_content_to_text(file_path: Path, file_type: str, workspace: Path) -> str:
    """Extract text content from different file types"""
    if file_type == 'pdf':
        return _extract_pdf_content(file_path)
    elif file_type == 'docx':
        return _extract_docx_content(file_path)
    elif file_type == 'excel':
        return _extract_excel_content(file_path)
    else:
        return file_path.read_text(encoding='utf-8')


def _extract_pdf_content(file_path: Path) -> str:
    """Extract text content from PDF files (reuses read_pdf logic)"""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    text += f"\n--- Page {page_num + 1} (Error) ---\nCould not extract text: {str(e)}\n"
            
            return text
    except ImportError:
        raise Exception("PDF processing not available - PyPDF2 not installed")


def _extract_docx_content(file_path: Path) -> str:
    """Extract text content from DOCX files (reuses read_docx logic)"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n\n"
        
        return text
    except ImportError:
        raise Exception("DOCX processing not available - python-docx not installed")


def _extract_excel_content(file_path: Path) -> str:
    """Extract text content from Excel files"""
    try:
        import openpyxl
        
        workbook = openpyxl.load_workbook(file_path)
        text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"\n--- Sheet: {sheet_name} ---\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None:
                        row_text.append(str(cell))
                if row_text:
                    text += "\t".join(row_text) + "\n"
        
        return text
    except ImportError:
        raise Exception("Excel processing not available - openpyxl not installed")


def write_file(file_path: str, content: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Write content to file"""
    try:
        # Validate content size (limit to 5MB)
        if len(content) > 5 * 1024 * 1024:
            content_size_mb = len(content) / (1024 * 1024)
            return {'success': False, 'output': f"❌ Content too large ({content_size_mb:.1f}MB) - maximum file size is 5MB"}
            
        resolved_path = resolve_path(file_path, workspace)
        resolved_path.parent.mkdir(parents=True, exist_ok=True)
        resolved_path.write_text(content, encoding='utf-8')
        return {
            'success': True,
            'output': f"✅ File written: {file_path} ({len(content)} chars)"
        }
    except PermissionError:
        return {'success': False, 'output': f"❌ Permission denied writing '{file_path}' - insufficient privileges to write to this location"}
    except IsADirectoryError:
        return {'success': False, 'output': f"❌ Cannot write '{file_path}' - this path is a directory, not a file"}
    except OSError as e:
        if "No space left on device" in str(e):
            return {'success': False, 'output': f"❌ Cannot write '{file_path}' - disk full, no space available"}
        elif "Read-only file system" in str(e):
            return {'success': False, 'output': f"❌ Cannot write '{file_path}' - file system is read-only"}
        else:
            return {'success': False, 'output': f"❌ System error writing '{file_path}': {str(e)} - file system issue"}
    except UnicodeEncodeError:
        return {'success': False, 'output': f"❌ Cannot write '{file_path}' - content contains characters that cannot be encoded as UTF-8"}
    except Exception as e:
        return {'success': False, 'output': f"❌ Error writing '{file_path}': {str(e)} - unable to write file content"}


def list_directory(path: str = ".", workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """List directory contents"""
    try:
        resolved_path = resolve_path(path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ Directory not found: '{path}' - check if the path exists"}
        if not resolved_path.is_dir():
            return {'success': False, 'output': f"❌ Not a directory: '{path}' - this is a file, not a directory"}
        
        items = []
        for item in resolved_path.iterdir():
            item_type = "📁" if item.is_dir() else "📄"
            size = f" ({item.stat().st_size} bytes)" if item.is_file() else ""
            items.append(f"{item_type} {item.name}{size}")
        
        return {
            'success': True,
            'output': f"📁 Directory: {path}\n" + "\n".join(sorted(items)),
            'items': [item.name for item in resolved_path.iterdir()]
        }
    except PermissionError:
        return {'success': False, 'output': f"❌ Permission denied accessing '{path}' - insufficient privileges to list this directory"}
    except NotADirectoryError:
        return {'success': False, 'output': f"❌ Not a directory: '{path}' - this is a file, not a directory"}
    except OSError as e:
        if "No such file or directory" in str(e):
            return {'success': False, 'output': f"❌ Directory not found: '{path}' - path does not exist"}
        else:
            return {'success': False, 'output': f"❌ System error accessing '{path}': {str(e)} - file system issue"}
    except Exception as e:
        return {'success': False, 'output': f"❌ Error listing directory '{path}': {str(e)} - unable to access directory contents"}


def find_files(pattern: str, path: str = ".", workspace: Path = DEFAULT_WORKSPACE, streaming_callback=None) -> Dict[str, Any]:
    """Search for files and directories with streaming progress"""
    try:
        if not pattern:
            return {'success': False, 'output': "❌ Pattern required"}
        
        # Secure path to workspace
        if path == ".":
            search_path = str(workspace)
        else:
            safe_path = Path(path).name.replace('..', '').replace('/', '_')
            search_path = str(workspace / safe_path)
        
        # Use streaming shell command for large file searches
        from .system_tools import shell_command
        command = f"find '{search_path}' -name '{pattern}' -type f"
        result = shell_command(command, workspace, streaming_callback)
        
        if result['success']:
            raw_output = result['raw_output'] if 'raw_output' in result else result['output']
            # Extract just the file paths from the formatted output
            lines = raw_output.split('\n')
            # Find the "Output:" line and get everything after it
            output_start = -1
            for i, line in enumerate(lines):
                if line.startswith('Output:'):
                    output_start = i + 1
                    break
            
            if output_start >= 0:
                file_lines = lines[output_start:]
            else:
                file_lines = lines
                
            files = [line.strip() for line in file_lines if line.strip() and not line.startswith('💻')]
            
            output = f"🔍 Found {len(files)} files matching '{pattern}':\n"
            for file_path in files[:20]:  # Show first 20 results
                relative_path = file_path.replace(str(workspace), '.')
                output += f"• {relative_path}\n"
            
            if len(files) > 20:
                output += f"... and {len(files) - 20} more files\n"
            
            return {
                'success': True,
                'output': output,
                'files_found': len(files),
                'pattern': pattern
            }
        else:
            return result
            
    except Exception as e:
        return {'success': False, 'output': f"❌ Find failed: {e}"}


def grep_search(pattern: str, file_path: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Search text in files"""
    try:
        if not pattern or not file_path:
            return {'success': False, 'output': "❌ Pattern and file path required"}
        
        # Secure file path
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ File not found: {file_path}"}
        
        command = ['grep', '-n', '-i', pattern, str(resolved_path)]
        result = execute_command(command, timeout=30)
        
        if result['success']:
            matches = result['output'].split('\n')
            match_count = len([m for m in matches if m.strip()])
            
            output = f"🔍 Found {match_count} matches for '{pattern}' in {file_path}:\n"
            for match in matches[:10]:  # Show first 10 matches
                if match.strip():
                    output += f"• {match}\n"
            
            if match_count > 10:
                output += f"... and {match_count - 10} more matches\n"
            
            return {
                'success': True,
                'output': output,
                'matches_found': match_count,
                'pattern': pattern
            }
        else:
            return {'success': False, 'output': f"❌ No matches found for '{pattern}'"}
            
    except Exception as e:
        return {'success': False, 'output': f"❌ Grep failed: {e}"}


def word_count(file_path: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Count words, lines, and characters in file"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        
        # Secure file path
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ File not found: {file_path}"}
        
        command = ['wc', str(resolved_path)]
        result = execute_command(command, timeout=30)
        
        if result['success']:
            # Parse wc output: lines words characters filename
            parts = result['output'].strip().split()
            if len(parts) >= 3:
                lines = parts[0]
                words = parts[1]
                chars = parts[2]
                
                output = f"📊 Word Count for {file_path}:\n"
                output += f"• Lines: {lines}\n"
                output += f"• Words: {words}\n"
                output += f"• Characters: {chars}"
                
                return {
                    'success': True,
                    'output': output,
                    'lines': int(lines),
                    'words': int(words),
                    'characters': int(chars)
                }
        
        return result
        
    except Exception as e:
        return {'success': False, 'output': f"❌ Word count failed: {e}"}


def file_head(file_path: str, lines: str = "10", workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Show first lines of file"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        
        # Validate lines parameter
        try:
            num_lines = int(lines)
            if num_lines < 1 or num_lines > 1000:
                num_lines = 10
        except:
            num_lines = 10
        
        # Secure file path
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ File not found: {file_path}"}
        
        command = ['head', '-n', str(num_lines), str(resolved_path)]
        result = execute_command(command, timeout=30)
        
        if result['success']:
            output = f"📄 First {num_lines} lines of {file_path}:\n{result['output']}"
            return {
                'success': True,
                'output': output,
                'lines_shown': num_lines
            }
        else:
            return result
            
    except Exception as e:
        return {'success': False, 'output': f"❌ Head failed: {e}"}


def file_tail(file_path: str, lines: str = "10", workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Show last lines of file"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        
        # Validate lines parameter
        try:
            num_lines = int(lines)
            if num_lines < 1 or num_lines > 1000:
                num_lines = 10
        except:
            num_lines = 10
        
        # Secure file path
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ File not found: {file_path}"}
        
        command = ['tail', '-n', str(num_lines), str(resolved_path)]
        result = execute_command(command, timeout=30)
        
        if result['success']:
            output = f"📄 Last {num_lines} lines of {file_path}:\n{result['output']}"
            return {
                'success': True,
                'output': output,
                'lines_shown': num_lines
            }
        else:
            return result
            
    except Exception as e:
        return {'success': False, 'output': f"❌ Tail failed: {e}"}


def file_info(file_path: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Get detailed file information"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        
        # Secure file path
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ File not found: {file_path}"}
        
        command = ['stat', str(resolved_path)]
        result = execute_command(command, timeout=30)
        
        if result['success']:
            output = f"📋 File Information for {file_path}:\n{result['output']}"
            return {
                'success': True,
                'output': output,
                'file_path': file_path
            }
        else:
            # Fallback using Python
            stat_info = resolved_path.stat()
            
            output = f"📋 File Information for {file_path}:\n"
            output += f"• Size: {stat_info.st_size} bytes\n"
            output += f"• Modified: {time.ctime(stat_info.st_mtime)}\n"
            output += f"• Permissions: {oct(stat_info.st_mode)[-3:]}"
            
            return {
                'success': True,
                'output': output,
                'size': stat_info.st_size,
                'modified': stat_info.st_mtime
            }
            
    except Exception as e:
        return {'success': False, 'output': f"❌ File info failed: {e}"}


# read_pdf function removed - use read_file() instead for all file types


def create_pdf(file_path: str, content: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Create PDF file from text content"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        if not content:
            return {'success': False, 'output': "❌ Content required"}
        
        # Validate content size (limit to 5MB)
        if len(content) > 5 * 1024 * 1024:
            content_size_mb = len(content) / (1024 * 1024)
            return {'success': False, 'output': f"❌ Content too large ({content_size_mb:.1f}MB) - maximum content size is 5MB"}
        
        resolved_path = resolve_path(file_path, workspace)
        resolved_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Import PDF creation library
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
        except ImportError:
            return {'success': False, 'output': "❌ PDF creation not available - reportlab not installed. Install with: pip install reportlab"}
        
        # Create PDF document
        doc = SimpleDocTemplate(str(resolved_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # Clean paragraph text for reportlab
                clean_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(clean_para, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(story)
        
        # Get file size after creation
        file_size = resolved_path.stat().st_size
        
        return {
            'success': True,
            'output': f"✅ PDF created: {file_path} ({file_size} bytes) - {len(paragraphs)} paragraphs",
            'file_path': str(resolved_path),
            'size': file_size,
            'paragraphs': len(paragraphs)
        }
        
    except PermissionError:
        return {'success': False, 'output': f"❌ Permission denied creating '{file_path}' - insufficient privileges to write to this location"}
    except Exception as e:
        return {'success': False, 'output': f"❌ Error creating PDF '{file_path}': {str(e)} - unable to create PDF file"}


# read_docx function removed - use read_file() instead for all file types


def create_docx(file_path: str, content: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Create DOCX file from text content"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        if not content:
            return {'success': False, 'output': "❌ Content required"}
        
        # Validate content size (limit to 5MB)
        if len(content) > 5 * 1024 * 1024:
            content_size_mb = len(content) / (1024 * 1024)
            return {'success': False, 'output': f"❌ Content too large ({content_size_mb:.1f}MB) - maximum content size is 5MB"}
        
        resolved_path = resolve_path(file_path, workspace)
        resolved_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Import DOCX library
        try:
            from docx import Document
        except ImportError:
            return {'success': False, 'output': "❌ DOCX creation not available - python-docx not installed. Install with: pip install python-docx"}
        
        # Create DOCX document
        doc = Document()
        
        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save document
        doc.save(str(resolved_path))
        
        # Get file size after creation
        file_size = resolved_path.stat().st_size
        
        return {
            'success': True,
            'output': f"✅ DOCX created: {file_path} ({file_size} bytes) - {len(paragraphs)} paragraphs",
            'file_path': str(resolved_path),
            'size': file_size,
            'paragraphs': len(paragraphs)
        }
        
    except PermissionError:
        return {'success': False, 'output': f"❌ Permission denied creating '{file_path}' - insufficient privileges to write to this location"}
    except Exception as e:
        return {'success': False, 'output': f"❌ Error creating DOCX '{file_path}': {str(e)} - unable to create DOCX file"}



def read_image_metadata(file_path: str, workspace: Path = DEFAULT_WORKSPACE) -> Dict[str, Any]:
    """Read metadata and information from image files"""
    try:
        if not file_path:
            return {'success': False, 'output': "❌ File path required"}
        
        resolved_path = resolve_path(file_path, workspace)
        if not resolved_path.exists():
            return {'success': False, 'output': f"❌ Image not found: '{file_path}' - check if the file exists or path is correct"}
        
        # Check file size (limit to 20MB for images)
        file_size = resolved_path.stat().st_size
        if file_size > 20 * 1024 * 1024:
            return {'success': False, 'output': f"❌ Image too large: '{file_path}' ({file_size // (1024*1024)}MB) - images over 20MB not supported"}
        
        # Import image library
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
        except ImportError:
            return {'success': False, 'output': "❌ Image processing not available - Pillow not installed. Install with: pip install Pillow"}
        
        # Open and analyze image
        with Image.open(resolved_path) as image:
            # Basic image info
            width, height = image.size
            format_name = image.format or "Unknown"
            mode = image.mode
            
            # Try to extract EXIF data
            exif_data = {}
            try:
                exif = image._getexif()
                if exif:
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_data[tag] = str(value)
            except:
                pass  # EXIF not available or readable
            
            # File size formatting
            if file_size < 1024:
                size_str = f"{file_size} bytes"
            elif file_size < 1024 * 1024:
                size_str = f"{file_size / 1024:.1f} KB"
            else:
                size_str = f"{file_size / (1024 * 1024):.1f} MB"
        
        output = f"🖼️ Image metadata for {file_path}:\n"
        output += f"• Format: {format_name}\n"
        output += f"• Dimensions: {width} x {height} pixels\n"
        output += f"• Color Mode: {mode}\n"
        output += f"• File Size: {size_str}\n"
        
        if exif_data:
            output += f"• EXIF Data: {len(exif_data)} tags found\n"
            # Show key EXIF data
            key_tags = ['DateTime', 'Make', 'Model', 'Software', 'Orientation']
            for tag in key_tags:
                if tag in exif_data:
                    output += f"  - {tag}: {exif_data[tag]}\n"
        else:
            output += "• EXIF Data: None found\n"
        
        return {
            'success': True,
            'output': output,
            'metadata': {
                'format': format_name,
                'width': width,
                'height': height,
                'mode': mode,
                'file_size': file_size,
                'file_size_formatted': size_str,
                'exif_data': exif_data
            }
        }
        
    except PermissionError:
        return {'success': False, 'output': f"❌ Permission denied reading '{file_path}' - insufficient privileges to access this file"}
    except Exception as e:
        return {'success': False, 'output': f"❌ Error reading image metadata '{file_path}': {str(e)} - unable to process image"}